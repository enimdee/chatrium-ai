#!/usr/bin/env python3
"""
Generate PRODUCT_BIBLE.docx from scratch using python-docx.
Run: python3 scripts/build_product_bible.py
Output: PRODUCT_BIBLE.docx (repo root)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(os.path.dirname(__file__)), "PRODUCT_BIBLE.docx")

# ── Colour palette ────────────────────────────────────────────────────────────
GOLD       = RGBColor(0xC5, 0xA5, 0x72)
DARK_BG    = RGBColor(0x1A, 0x1D, 0x21)
HEADING_FG = RGBColor(0x1A, 0x1D, 0x21)   # dark on gold header
BODY_FG    = RGBColor(0x2C, 0x2C, 0x2C)
MUTED_FG   = RGBColor(0x6B, 0x72, 0x80)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GOLD = RGBColor(0xFA, 0xF3, 0xE4)   # table header tint

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, color=None, size=10, align=None):
    para = cell.paragraphs[0]
    para.clear()
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.color.rgb = GOLD if level == 1 else DARK_BG
    run.font.size = Pt(16 if level == 1 else 13)
    return p

def add_body(text, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = italic
    run.font.color.rgb = color or BODY_FG
    return p

def add_bullet(text, sub=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if sub:
        p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_FG

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hdr.cells[i], GOLD)
        cell_text(hdr.cells[i], h, bold=True, color=HEADING_FG, size=10)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_cell_bg(cell, LIGHT_GOLD)
            cell_text(cell, val, size=10, color=BODY_FG)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table

def hr():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("─" * 72)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

# ═══════════════════════════════════════════════════════════════════════════════
#  COVER
# ═══════════════════════════════════════════════════════════════════════════════

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(40)
title_p.paragraph_format.space_after  = Pt(6)
r = title_p.add_run("PRODUCT BIBLE")
r.font.size = Pt(28)
r.font.bold = True
r.font.color.rgb = GOLD

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_after = Pt(4)
r2 = sub_p.add_run("Wokeflow AI Communication Assistant")
r2.font.size = Pt(16)
r2.font.color.rgb = BODY_FG

tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_p.paragraph_format.space_after = Pt(32)
r3 = tag_p.add_run("สำหรับอุตสาหกรรมการบริการ (Hospitality)")
r3.font.size = Pt(12)
r3.font.color.rgb = MUTED_FG
r3.font.italic = True

meta = [
    ("เวอร์ชัน", "1.0"),
    ("วันที่", "21 เมษายน 2569"),
    ("เจ้าของ Product", "Sakchai Nimdee — Wokeflow"),
    ("สถานะ", "Pilot Production (Chatrium Rawai Phuket)"),
]
meta_t = doc.add_table(rows=len(meta), cols=2)
meta_t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(meta):
    cell_text(meta_t.rows[i].cells[0], k, bold=True, color=MUTED_FG, size=10)
    cell_text(meta_t.rows[i].cells[1], v, color=BODY_FG, size=10)
    meta_t.rows[i].cells[0].width = Cm(5)
    meta_t.rows[i].cells[1].width = Cm(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. ONE-LINER
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("1. ONE-LINER")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
p.paragraph_format.left_indent = Inches(0.3)
r = p.add_run(
    "Wokeflow AI ช่วยให้พนักงานโรงแรมเขียน email ภาษาอังกฤษระดับมืออาชีพที่ตรงกับ "
    "brand voice ของโรงแรม ภายใน 10 วินาที โดยพิมพ์เพียงวัตถุประสงค์เป็นภาษาไทย"
)
r.font.size = Pt(12)
r.font.italic = True
r.font.color.rgb = BODY_FG
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("2. ปัญหาที่แก้ไข (Problem Statement)")
add_body(
    "โรงแรมระดับ 4–5 ดาวในไทยมีพนักงานที่ต้องเขียน email ภาษาอังกฤษวันละ 20–50 ฉบับ "
    "ได้แก่ Front Office, Reservations, Sales, Guest Relations และ GM "
    "ปัญหาเกิดขึ้น 3 ระดับ:"
)

add_table(
    ["ระดับ", "ปัญหา", "ผลกระทบ"],
    [
        ["พนักงาน", "ภาษาอังกฤษไม่มั่นใจ ใช้เวลาเขียน 10–20 นาที/ฉบับ", "ทำงานอื่นได้น้อยลง / เครียด"],
        ["Manager", "draft ที่พนักงานส่งมักไม่ตรง brand voice", "แก้ไขซ้ำ เสียเวลา GM"],
        ["MARCOM/HQ", "ไม่รู้ว่า email ที่ส่งออกทุกวัน on-brand หรือไม่", "ความสม่ำเสมอของ brand เสีย"],
    ],
    col_widths=[3, 7, 6],
)

p = doc.add_paragraph()
run = p.add_run("ทำไม solution เดิมไม่พอ:")
run.font.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = BODY_FG
p.paragraph_format.space_after = Pt(4)

add_table(
    ["วิธีเดิม", "ข้อเสีย"],
    [
        ["ChatGPT / Gemini (free)", "ไม่รู้จัก brand voice ของโรงแรม ต้องพิมพ์ context ทุกครั้ง"],
        ["Grammarly", "แก้ grammar เท่านั้น ไม่ generate draft ให้"],
        ["Template Word/Excel", "ไม่ยืดหยุ่น ยังต้องเขียนเนื้อหาหลักเอง"],
        ["จ้าง native speaker review", "ช้า แพง ไม่ scalable"],
    ],
    col_widths=[5, 11],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. SOLUTION
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("3. Solution")
add_body(
    "Wokeflow AI คือ web application ที่ฝัง brand voice ของโรงแรมไว้ใน AI model ถาวร "
    "พนักงานพิมพ์แค่วัตถุประสงค์ (ภาษาไทยหรืออังกฤษ) แล้ว AI generate email สำเร็จรูปออกมาพร้อมใช้ "
    "ผ่าน Outlook ได้ใน 1 คลิก"
)

p = doc.add_paragraph()
r = p.add_run("กระบวนการ (User Flow):")
r.font.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = BODY_FG
p.paragraph_format.space_after = Pt(4)

steps = [
    "1. พนักงานเปิด ai.[hotel].wokeflow.net",
    "2. เลือก Template (หรือพิมพ์ objective เอง)",
    "3. AI generate draft ภายใน ~10 วินาที",
    "4. ตรวจ → Refine ถ้าต้องการ (Shorter / Warmer / More formal / Custom)",
    "5. กด \"Open in Outlook\" → 1 คลิก → เลือก recipient → ส่ง",
]
for s in steps:
    add_bullet(s)

doc.add_paragraph()
add_table(
    ["ขั้นตอน", "เดิม", "ใหม่"],
    [
        ["คิด/เขียน draft", "10–20 นาที", "พิมพ์ objective ~2 นาที"],
        ["แก้ไขตาม brand voice", "5–10 นาที (GM review)", "อัตโนมัติ — 0 นาที"],
        ["Copy → Outlook", "2–3 นาที (7 steps)", "1 คลิก (~10 วินาที)"],
        ["รวม", "~20–30 นาที/ฉบับ", "~3–5 นาที/ฉบับ"],
    ],
    col_widths=[5, 5, 6],
)
add_body(
    "โรงแรมที่มีพนักงานส่ง email 30 ฉบับ/วัน → ประหยัดได้ 8–12 ชั่วโมง/วัน",
    italic=True, color=GOLD
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. TARGET MARKET
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("4. กลุ่มเป้าหมาย (Target Market)")

add_table(
    ["Segment", "จำนวน", "Notes"],
    [
        ["โรงแรมระดับ 4–5 ดาว (กทม + รีสอร์ท)", "~800 แห่ง", "ตลาดหลัก"],
        ["โรงแรม boutique 3 ดาวขึ้นไป", "~2,000 แห่ง", "Secondary"],
        ["Hotel chains (Chatrium, Centara, Dusit)", "~50 groups", "Enterprise"],
        ["โรงแรมในอาเซียน (เฟส 2+)", "~5,000 แห่ง", "ขยาย"],
    ],
    col_widths=[7, 4, 5],
)

p = doc.add_paragraph()
r = p.add_run("Beachhead: Chatrium Hotels & Residences")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BODY_FG
p.paragraph_format.space_after = Pt(4)

beachhead = [
    "9 properties ในไทย (~2,500 rooms)",
    "Pilot: Chatrium Rawai Phuket (~30 users)",
    "Rollout target: Grand Chatrium Bangkok, Maitria Sukhumvit (~200 users)",
    "Case study → ขาย properties อื่น",
]
for b in beachhead:
    add_bullet(b)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("User Personas")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BODY_FG

add_table(
    ["Persona", "บทบาท", "Pain Point", "Goal"],
    [
        ["Pim (28)", "Front Office Staff", "ภาษาอังกฤษไม่มั่นใจ เขียน 30–40 emails/วัน", "ส่ง email ถูกต้อง รวดเร็ว"],
        ["Khun Nok (45)", "General Manager", "review draft ทีม 5–10 ฉบับ/วัน", "ลด review time + เห็น analytics"],
        ["MARCOM Manager", "HQ Brand Keeper", "ไม่รู้ว่า email on-brand หรือไม่", "Central control + report CEO"],
    ],
    col_widths=[3.5, 3.5, 5.5, 4.5],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. FEATURES
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("5. Features หลัก")

p = doc.add_paragraph()
r = p.add_run("✅ Launched (v1.0 — เมษายน 2569)")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

add_table(
    ["Feature", "คำอธิบาย"],
    [
        ["AI Email Draft", "Generate email ภาษาอังกฤษจาก objective ภาษาไทย/อังกฤษ"],
        ["Brand Voice Enforcement", "ฝัง brand rules ถาวรใน AI — พนักงานเปลี่ยนไม่ได้"],
        ["Magic-link Auth", "Login ผ่าน email link ไม่ต้องจำ password"],
        ["Draft History", "บันทึกทุก draft — click กลับมาแก้ต่อได้"],
        ["Outlook Deep-link", "1 คลิกเปิด Outlook พร้อม subject + body"],
        ["Regenerate / Refine", "Shorter / Warmer / More formal / Custom instruction"],
        ["Template Library", "8 templates สำเร็จรูปสำหรับ scenarios ที่พบบ่อย"],
        ["Multi-language Output", "Generate email ภาษา EN / 中文 / 日本語 / 한국어 / ไทย"],
        ["QC Auto-check", "ตรวจ 5 จุด: no em-dash, no slang, CTA, loyalty mention, length"],
        ["Cost Ceiling", "จำกัด cost/เดือน ป้องกัน AI cost เกินงบ"],
        ["Admin Dashboard", "Brand voice editor, user management, usage analytics"],
        ["Multi-property", "รองรับหลาย property ใน 1 instance"],
    ],
    col_widths=[5, 11],
)

p = doc.add_paragraph()
r = p.add_run("🔜 Roadmap (เฟส 2 — Q3 2569)")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

add_table(
    ["Priority", "Feature", "Impact"],
    [
        ["High", "Brand Voice Score", "คะแนน 0–100 ต่อ draft — MARCOM มี data"],
        ["High", "Analytics for Managers", "draft/user/week, avg score, time saved"],
        ["Medium", "Recipient Memory", "จำ context ลูกค้า repeat guest"],
        ["Medium", "Team Review Workflow", "FO draft → GM approve → send"],
        ["Medium", "Mobile-responsive", "ใช้ได้บนมือถือ"],
    ],
    col_widths=[3, 5, 8],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. BUSINESS MODEL
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("6. Business Model")

add_table(
    ["Plan", "ราคา/เดือน", "Users", "Features"],
    [
        ["Starter", "฿1,990", "≤10 users", "Core features, 1 property"],
        ["Professional", "฿3,990", "≤50 users", "All features, 3 properties, analytics"],
        ["Enterprise", "Custom", "Unlimited", "SSO, SLA, custom brand voice, HQ dashboard"],
    ],
    col_widths=[4, 3.5, 3.5, 6],
)
add_body("* ตัวเลข pricing เป็น draft — ยืนยันกับ pilot feedback ก่อน go-live", italic=True, color=MUTED_FG)

p = doc.add_paragraph()
r = p.add_run("Unit Economics (ประมาณการต่อ instance/เดือน)")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BODY_FG

add_table(
    ["รายการ", "ต่อ instance/เดือน"],
    [
        ["Claude API cost (500 drafts × ฿0.15)", "~฿75"],
        ["VPS infra cost (shared)", "~฿200"],
        ["COGS รวม", "~฿275"],
        ["Revenue (Professional)", "฿3,990"],
        ["Gross Margin", "~86%"],
    ],
    col_widths=[9, 4],
)

p = doc.add_paragraph()
r = p.add_run("Revenue Targets")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BODY_FG

add_table(
    ["Period", "Customers", "MRR", "ARR"],
    [
        ["Year 1", "10 hotels", "฿39,900", "~฿480,000"],
        ["Year 2", "50 hotels", "฿199,500", "~฿2,400,000"],
    ],
    col_widths=[3.5, 3.5, 4, 5.5],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. TECH STACK
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("7. เทคโนโลยี (Technical Stack)")

add_table(
    ["ชั้น", "เทคโนโลยี", "เหตุผล"],
    [
        ["Frontend", "Next.js 15, React 19, Tailwind 4", "Modern, fast, SSR"],
        ["AI", "Claude Sonnet (Anthropic) + Prompt Caching", "คุณภาพสูง, cost ต่ำ 90% ด้วย caching"],
        ["Auth", "Custom magic-link (JWT + HMAC-SHA256)", "ไม่ต้องพึ่ง third-party, Edge-compatible"],
        ["Data Phase 1", "File-based JSON", "Deploy ง่าย, scale ถึง 200 users/instance"],
        ["Data Phase 2", "MySQL 8 + Drizzle ORM", "Multi-tenant ready"],
        ["Deploy", "Docker Compose + Traefik + Hostinger VPS", "Per-customer isolation"],
        ["CI/CD", "GitHub Actions → ghcr.io → VPS", "Push-to-deploy ใน 3 นาที"],
    ],
    col_widths=[3, 6, 7],
)

add_body("VPS Capacity: รองรับ 5–6 customers ต่อ VPS (KVM2 / 8GB RAM)  |  Scale: KVM4 (16GB) → 12–14 customers")
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. COMPETITIVE ADVANTAGE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("8. Competitive Advantage")

add_table(
    ["ข้อได้เปรียบ", "คำอธิบาย"],
    [
        ["Brand Voice Lock-in", "ฝัง brand rules ใน system prompt — พนักงานเปลี่ยนไม่ได้ output สม่ำเสมอ 100%"],
        ["Hospitality-specific", "Templates, task types, QC rules ออกแบบมาสำหรับโรงแรมโดยเฉพาะ"],
        ["Thai-first Input", "Staff พิมพ์ภาษาไทยได้ ไม่ต้องแปลก่อน — ลด friction หลัก"],
        ["Multi-language Output", "สร้าง email ภาษาจีน/ญี่ปุ่น/เกาหลีได้ใน 1 คลิก — Critical สำหรับ APAC"],
        ["Outlook Integration", "1 คลิกเปิด Outlook พร้อม pre-fill — ลด workflow จาก 7 steps เหลือ 2"],
        ["Per-instance Isolation", "Data ของแต่ละโรงแรมแยกกันสมบูรณ์ — ตอบโจทย์ enterprise ด้าน security"],
        ["Cost Predictability", "Hard ceiling รายเดือน ป้องกัน overspend AI cost"],
        ["Deploy in 10 min", "Onboard customer ใหม่ได้ใน 10 นาที ด้วย script เดียว"],
    ],
    col_widths=[5, 11],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. NON-GOALS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("9. Constraints & Non-goals")

p = doc.add_paragraph()
r = p.add_run("สิ่งที่ตั้งใจไม่ทำ (Non-goals)")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BODY_FG

non_goals = [
    "ส่ง email อัตโนมัติ — พนักงานต้อง review ก่อนส่งเสมอ (human-in-the-loop)",
    "Multi-tenant DB — แต่ละ customer มี DB และ container ของตัวเอง",
    "Voice input / Image generation / File upload",
    "Fine-tune custom model — Claude ดีพอแล้ว, prompt engineering เพียงพอ",
    "Real-time collaboration (Google Docs style)",
    "CRM / PMS replacement — เป็นแค่ email drafting tool",
]
for ng in non_goals:
    add_bullet(f"✗  {ng}")
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. SUCCESS METRICS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("10. Metrics ความสำเร็จ")

p = doc.add_paragraph()
r = p.add_run("Pilot KPIs (Chatrium Rawai — 3 เดือนแรก)")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BODY_FG

add_table(
    ["Metric", "Target"],
    [
        ["DAU / Registered Users", "≥ 50% (15/30 users ใช้ทุกวัน)"],
        ["Drafts per active user/day", "≥ 5 drafts"],
        ["Time-to-draft (median)", "≤ 90 วินาที"],
        ["GM review rejection rate", "ลดลง ≥ 50% vs baseline"],
        ["NPS (staff survey)", "≥ 40"],
        ["Cost per draft", "≤ ฿0.20"],
    ],
    col_widths=[8, 8],
)

p = doc.add_paragraph()
r = p.add_run("Business KPIs (Year 1)")
r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BODY_FG

add_table(
    ["Metric", "Target"],
    [
        ["MRR", "฿40,000 (10 customers)"],
        ["Customer churn", "≤ 5%/เดือน"],
        ["Gross Margin", "≥ 80%"],
        ["CAC", "≤ ฿10,000"],
        ["LTV", "≥ ฿120,000 (2 ปี × Professional)"],
    ],
    col_widths=[8, 8],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. RISKS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("11. ความเสี่ยงและการรับมือ")

add_table(
    ["ความเสี่ยง", "โอกาส", "ผลกระทบ", "วิธีรับมือ"],
    [
        ["Staff ไม่ยอม adopt", "Medium", "High", "Training + template ลด friction + quick win demo"],
        ["AI hallucinate ข้อมูลผิด", "Low", "High", "Staff review ก่อนส่งเสมอ, QC auto-check"],
        ["Anthropic เพิ่มราคา API", "Low", "Medium", "Cost ceiling + switching to OpenAI/Gemini ready"],
        ["โรงแรมเปลี่ยน brand voice", "Low", "Medium", "Admin self-serve editor — MARCOM แก้ได้เอง"],
        ["Data breach", "Very Low", "High", "Per-instance isolation, httpOnly cookie, no PII in AI"],
        ["Competitor ทำเหมือนกัน", "Medium", "Medium", "Brand voice lock-in + hospitality UX = moat"],
    ],
    col_widths=[5, 2.5, 2.5, 6],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. OPEN DECISIONS
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("12. Open Decisions")

add_table(
    ["ID", "เรื่อง", "ผู้รับผิดชอบ", "Deadline"],
    [
        ["D-1", "Pricing final (pilot feedback)", "Sakchai", "หลัง pilot เดือน 1"],
        ["D-2", "Anthropic billing ownership (corporate card)", "Sakchai", "ก่อน go-live"],
        ["D-3", "Brand voice MARCOM sign-off (Chatrium)", "Richard Mehr", "เดือน 1"],
        ["D-4", "Pilot user list 30 คน", "Richard Mehr", "เดือน 1"],
        ["D-5", "Subdomain final (ai.chatrium.com vs ai.wokeflow.net)", "Sakchai + IT", "ก่อน staff training"],
    ],
    col_widths=[1.5, 6, 4, 4],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  13. PEOPLE
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("13. People")

add_table(
    ["บทบาท", "คน", "Contact"],
    [
        ["Product / Engineering", "Sakchai Nimdee (Wokeflow)", "sakchai.nim@chatrium.com"],
        ["Brand Voice", "MARCOM via Sakchai", "—"],
        ["Pilot Customer", "Richard A. Mehr", "richard.meh@chatrium.com"],
        ["VPS / Infra", "Anuwat", "anuwat.wat@chatrium.com"],
    ],
    col_widths=[5, 5, 6],
)
hr()

# ═══════════════════════════════════════════════════════════════════════════════
#  14. CHANGELOG
# ═══════════════════════════════════════════════════════════════════════════════
add_heading("14. Changelog")

add_table(
    ["วันที่", "เวอร์ชัน", "รายการ"],
    [
        ["21 เม.ย. 2569", "1.0", "Initial Product Bible — post Tier A feature complete"],
        ["—", "0.9", "Sprint 1 (Auth + Cost tracking + Wokeflow branding)"],
        ["—", "0.1", "Phase 1 scaffold (compose + generate API)"],
    ],
    col_widths=[4, 3, 9],
)

# Footer note
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(20)
r = p.add_run(
    "Document นี้เป็น living document — อัพเดทเมื่อ feature, pricing หรือ strategy เปลี่ยน\n"
    "เจ้าของ: Sakchai Nimdee / Wokeflow — อย่า distribute โดยไม่ได้รับอนุญาต"
)
r.font.size = Pt(9)
r.font.italic = True
r.font.color.rgb = MUTED_FG

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"✔ Saved: {OUT}")
